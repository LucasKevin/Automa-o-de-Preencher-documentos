from tkinter import *
import os
from docx import Document
from docx2pdf import convert
from tkinter import filedialog
from tkinter import messagebox


def preencher_documento():
    # Carregar o modelo do documento
    doc = Document('declaração isencao DANFe - DHL.docx')
    
    # Dicionário com os campos do documento e suas respectivas entradas
    campos = {
        '<<AWB>>': awb_entry.get(),
        '<<VALOR>>': valor_entry.get(),
        '<<DESCRICAO_NOTEBOOK>>': descricao_notebook_entry.get(),
        '<<VALOR_TOTAL>>': valor_total_entry.get(),
    }

    # Iterar sobre o dicionário e substituir os campos no documento
    for campo, entrada in campos.items():
        if entrada:
            for p in doc.paragraphs:
                if campo in p.text:
                    p.text = p.text.replace(campo, entrada)
    
    # Iterar sobre os parágrafos para tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for campo, entrada in campos.items():
                        if campo in paragraph.text:
                            paragraph.text = paragraph.text.replace(campo, entrada)

    # Salvar o doc preenchido como docx temporário
    temp_docx = 'Declaração isenção DANFe - DHL preenchido.docx'
    doc.save(temp_docx)

    # Converter o documento docx para pdf
    convert(temp_docx)  # Isso criará um arquivo 'documento_preenchido.pdf'

    # Salvar o documento preenchido
    doc.save('Declaração isenção DANFe - DHL preenchido.docx')

    # Limpar os campos do formulário
    limpar_campos()

def preencher_e_salvar():
    # Verificar se todos os campos foram preenchidos
    campos = [
        awb_entry.get(),
        valor_entry.get(),
        descricao_notebook_entry.get(),
        valor_total_entry.get(),
    ]
    
    if all(campos):  # Se todos os campos foram preenchidos
        preencher_documento()

        # Renomear o arquivo PDF
        novo_nome = 'Declaração isenção DANFe - DHL preenchido.docx'

        # Mover o arquivo PDF para o diretório desejado
        # Certifique-se de especificar o diretório correto onde você deseja salvar o arquivo PDF
        novo_caminho = filedialog.asksaveasfilename(defaultextension=".pdf", initialfile=novo_nome, filetypes=[("PDF files", "*.pdf")])
        if novo_caminho:
            os.rename('Declaração isenção DANFe - DHL preenchido.pdf', novo_caminho)
    else:  # Se algum campo estiver vazio
        messagebox.showerror("Erro", "Por favor, preencha todos os campos antes de submeter.")


def limpar_campos():
    awb_entry.delete(0, 'end')
    valor_entry.delete(0, 'end')
    descricao_notebook_entry.delete(0, 'end')
    valor_total_entry.delete(0, 'end')

# Criar a janela principal
root = Tk()
root.title("Preencher Documento")

# Criar rótulos e campos para o formulário
Label(root, text="AWB:").grid(row=0, column=0, sticky=W)
awb_entry = Entry(root)
awb_entry.grid(row=0, column=1)

Label(root, text="Valor:").grid(row=1, column=0, sticky=W)
valor_entry = Entry(root)
valor_entry.grid(row=1, column=1)

Label(root, text="Descrição:").grid(row=2, column=0, sticky=W)
descricao_notebook_entry = Entry(root)
descricao_notebook_entry.grid(row=2, column=1)

Label(root, text="Valor Total:").grid(row=3, column=0, sticky=W)
valor_total_entry = Entry(root)
valor_total_entry.grid(row=3, column=1)


# Botão de envio
Button(root, text="Preencher e Salvar", command=preencher_e_salvar).grid(row=10, column=0, columnspan=2)

# Iniciar o loop principal da interface gráfica
root.mainloop()
