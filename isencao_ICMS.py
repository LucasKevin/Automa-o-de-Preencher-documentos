from tkinter import *
import os
from docx import Document
from docx2pdf import convert
from tkinter import filedialog
from tkinter import messagebox


def preencher_documento():
    # Carregar o modelo do documento
    doc = Document('declaracao_ICMS.docx')
    
    # Dicionário com os campos do documento e suas respectivas entradas
    campos = {
        '<<MODELO>>': modelo_entry.get(),
        '<<SERVICETAG>>': servicetag_entry.get(),
        '<<VALOR>>': valor_entry.get(),
        '<<NOME>>': nome_entry.get(),
        '<<RG>>': rg_entry.get(),
        '<<CPF>>': cpf_entry.get(),
        '<<ENDERECO_COMPLETO>>': endereco_entry.get(),
        '<<LOCAL_E_DATA>>': local_e_data_entry.get(),
    }

    # Iterar sobre o dicionário e substituir os campos no documento
    for campo, entrada in campos.items():
        if entrada:
            for p in doc.paragraphs:
                if campo in p.text:
                    p.text = p.text.replace(campo, entrada)
    
    # Iterar sobre os parágrafos para tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for campo, entrada in campos.items():
                        if campo in paragraph.text:
                            paragraph.text = paragraph.text.replace(campo, entrada)

    # Salvar o doc preenchido como docx temporário
    temp_docx = 'declaracao_ICMS - preenchido.docx'
    doc.save(temp_docx)

    # Converter o documento docx para pdf
    convert(temp_docx)  # Isso criará um arquivo 'documento_preenchido.pdf'

    # Salvar o documento preenchido
    doc.save('declaracao_ICMS - preenchido.docx')

    # Limpar os campos do formulário
    limpar_campos()

def preencher_e_salvar():
    # Verificar se todos os campos foram preenchidos
    campos = [
        modelo_entry.get(),
        servicetag_entry.get(),
        valor_entry.get(),
        nome_entry.get(),
        endereco_entry.get(),
        rg_entry.get(),
        cpf_entry.get(),
        endereco_entry.get(),
        local_e_data_entry.get(),
    ]
    
    if all(campos):  # Se todos os campos foram preenchidos
        preencher_documento()

        # Renomear o arquivo PDF
        novo_nome = 'declaracao_ICMS - preenchido.docx'

        # Mover o arquivo PDF para o diretório desejado
        # Certifique-se de especificar o diretório correto onde você deseja salvar o arquivo PDF
        novo_caminho = filedialog.asksaveasfilename(defaultextension=".pdf", initialfile=novo_nome, filetypes=[("PDF files", "*.pdf")])
        if novo_caminho:
            os.rename('declaracao_ICMS - preenchido.pdf', novo_caminho)
    else:  # Se algum campo estiver vazio
        messagebox.showerror("Erro", "Por favor, preencha todos os campos antes de submeter.")


def limpar_campos():
    modelo_entry.delete(0, 'end')
    servicetag_entry.delete(0, 'end')
    valor_entry.delete(0, 'end')
    nome_entry.delete(0, 'end')
    rg_entry.delete(0, 'end')
    cpf_entry.delete(0, 'end')
    endereco_entry.delete(0, 'end')
    local_e_data_entry.delete(0, 'end')

# Criar a janela principal
root = Tk()
root.title("Preencher Documento")

# Criar rótulos e campos para o formulário
Label(root, text="Modelo:").grid(row=0, column=0, sticky=W)
modelo_entry = Entry(root)
modelo_entry.grid(row=0, column=1)

Label(root, text="Service Tag:").grid(row=1, column=0, sticky=W)
servicetag_entry = Entry(root)
servicetag_entry.grid(row=1, column=1)

Label(root, text="Valor:").grid(row=2, column=0, sticky=W)
valor_entry = Entry(root)
valor_entry.grid(row=2, column=1)

Label(root, text="Nome:").grid(row=3, column=0, sticky=W)
nome_entry = Entry(root)
nome_entry.grid(row=3, column=1)

Label(root, text="Rg:").grid(row=4, column=0, sticky=W)
rg_entry = Entry(root)
rg_entry.grid(row=4, column=1)

Label(root, text="CPF:").grid(row=5, column=0, sticky=W)
cpf_entry = Entry(root)
cpf_entry.grid(row=5, column=1)

Label(root, text="Endereço Completo:").grid(row=6, column=0, sticky=W)
endereco_entry = Entry(root)
endereco_entry.grid(row=6, column=1)

Label(root, text="Local e Data:").grid(row=7, column=0, sticky=W)
local_e_data_entry = Entry(root)
local_e_data_entry.grid(row=7, column=1)


# Botão de envio
Button(root, text="Preencher e Salvar", command=preencher_e_salvar).grid(row=10, column=0, columnspan=2)

# Iniciar o loop principal da interface gráfica
root.mainloop()