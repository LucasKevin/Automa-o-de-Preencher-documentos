from tkinter import *
import os
from docx import Document
from docx2pdf import convert
from tkinter import messagebox
from tkinter import filedialog

def preencher_documento():
    # Carregar o modelo do documento
    doc = Document('declaracao_NF.docx')
    
    # Dicionário com os campos do documento e suas respectivas entradas
    campos = {
        '<<AME>>': ame_entry.get(),
        '<<AWB>>': awb_entry.get(),
        '<<REMETENTE>>': remetente_entry.get(),
        '<<CPF_CNPJ>>': cpf_cnpj_entry.get(),
        '<<ENDERECO>>': endereco_entry.get(),
        '<<CIDADE>>': cidade_entry.get(),
        '<<DESTINATARIO>>': destinatario_entry.get(),
        '<<CPF_CNPJ2>>': cpf_cnpj2_entry.get(),
        '<<ENDERECO2>>': endereco2_entry.get(),
        '<<CIDADE>>': cidade_entry.get(),
        '<<DESCRICAO>>': descricao_entry.get(),
        '<<VALORUN>>': valorun_entry.get(),
        '<<VALOR_TOT>>': valor_tot_entry.get(),
        '<<VALOR_FINAL>>': valor_final_entry.get(),
        '<<LOCAL_E_DATA>>': local_e_data_entry.get(),
    }

    # Iterar sobre o dicionário e substituir os campos no documento
    for campo, entrada in campos.items():
        if entrada:
            for p in doc.paragraphs:
                if campo in p.text:
                    p.text = p.text.replace(campo, entrada)
    
    # Iterar sobre os parágrafos para tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for campo, entrada in campos.items():
                        if campo in paragraph.text:
                            paragraph.text = paragraph.text.replace(campo, entrada)

    # Salvar o doc preenchido como docx temporário
    temp_docx = 'declaracao_NF - preenchido.docx'
    doc.save(temp_docx)

    # Converter o documento docx para pdf
    convert(temp_docx)  # Isso criará um arquivo 'documento_preenchido.pdf'

    # Salvar o documento preenchido
    doc.save('declaracao_NF - preenchido.docx')

    # Limpar os campos do formulário
    limpar_campos()

def preencher_e_salvar():
    # Verificar se todos os campos foram preenchidos
    campos = [
        ame_entry.get(),
        awb_entry.get(),
        remetente_entry.get(),
        cpf_cnpj_entry.get(),
        endereco_entry.get(),
        cidade_entry.get(),
        destinatario_entry.get(),
        cpf_cnpj2_entry.get(),
        endereco2_entry.get(),
        cidade2_entry.get(),
        descricao_entry.get(),
        valorun_entry.get(),
        valor_tot_entry.get(),
        valor_final_entry.get(),
        local_e_data_entry.get(),
    ]
    
    if all(campos):  # Se todos os campos foram preenchidos
        preencher_documento()

        # Renomear o arquivo PDF
        novo_nome = 'declaracao_NF - preenchido.docx'

        # Mover o arquivo PDF para o diretório desejado
        # Certifique-se de especificar o diretório correto onde você deseja salvar o arquivo PDF
        novo_caminho = filedialog.asksaveasfilename(defaultextension=".pdf", initialfile=novo_nome, filetypes=[("PDF files", "*.pdf")])
        if novo_caminho:
            os.rename('declaracao_NF - preenchido.pdf', novo_caminho)
    else:  # Se algum campo estiver vazio
        messagebox.showerror("Erro", "Por favor, preencha todos os campos antes de submeter.")

def limpar_campos():
    ame_entry.delete(0, 'end')
    awb_entry.delete(0, 'end')
    remetente_entry.delete(0, 'end')
    cpf_cnpj_entry.delete(0, 'end')
    endereco_entry.delete(0, 'end')
    cidade_entry.delete(0, 'end')
    destinatario_entry.delete(0, 'end')
    cpf_cnpj2_entry.delete(0, 'end')
    endereco2_entry.delete(0, 'end')
    cidade2_entry.delete(0, 'end')
    descricao_entry.delete(0, 'end')
    valorun_entry.delete(0, 'end')
    valor_tot_entry.delete(0, 'end')
    valor_final_entry.delete(0, 'end')
    local_e_data_entry.delete(0, 'end')

# Criar a janela principal
root = Tk()
root.title("Preencher Documento")

# Criar rótulos e campos para o formulário
Label(root, text="AME:").grid(row=0, column=0, sticky=W)
ame_entry = Entry(root)
ame_entry.grid(row=0, column=1)

Label(root, text="AWB:").grid(row=1, column=0, sticky=W)
awb_entry = Entry(root)
awb_entry.grid(row=1, column=1)

Label(root, text="Remetente:").grid(row=2, column=0, sticky=W)
remetente_entry = Entry(root)
remetente_entry.grid(row=2, column=1)

Label(root, text="CPF ou CNPJ:").grid(row=3, column=0, sticky=W)
cpf_cnpj_entry = Entry(root)
cpf_cnpj_entry.grid(row=3, column=1)

Label(root, text="Endereço:").grid(row=4, column=0, sticky=W)
endereco_entry = Entry(root)
endereco_entry.grid(row=4, column=1)

Label(root, text="Cidade:").grid(row=5, column=0, sticky=W)
cidade_entry = Entry(root)
cidade_entry.grid(row=5, column=1)

Label(root, text="Destinatário:").grid(row=6, column=0, sticky=W)
destinatario_entry = Entry(root)
destinatario_entry.grid(row=6, column=1)

Label(root, text="CPF ou CNPJ:").grid(row=7, column=0, sticky=W)
cpf_cnpj2_entry = Entry(root)
cpf_cnpj2_entry.grid(row=7, column=1)

Label(root, text="Endereço:").grid(row=8, column=0, sticky=W)
endereco2_entry = Entry(root)
endereco2_entry.grid(row=8, column=1)

Label(root, text="Cidade:").grid(row=9, column=0, sticky=W)
cidade2_entry = Entry(root)
cidade2_entry.grid(row=9, column=1)

Label(root, text="Descrição Notebook:").grid(row=10, column=0, sticky=W)
descricao_entry = Entry(root)
descricao_entry.grid(row=10, column=1)

Label(root, text="Valor unitário:").grid(row=11, column=0, sticky=W)
valorun_entry = Entry(root)
valorun_entry.grid(row=11, column=1)

Label(root, text="Valor total:").grid(row=12, column=0, sticky=W)
valor_tot_entry = Entry(root)
valor_tot_entry.grid(row=12, column=1)

Label(root, text="Valor final:").grid(row=13, column=0, sticky=W)
valor_final_entry = Entry(root)
valor_final_entry.grid(row=13, column=1)

Label(root, text="local e Data:").grid(row=14, column=0, sticky=W)
local_e_data_entry = Entry(root)
local_e_data_entry.grid(row=14, column=1)


# Botão de envio
Button(root, text="Preencher e Salvar", command=preencher_e_salvar).grid(row=15, column=0, columnspan=2)

# Iniciar o loop principal da interface gráfica
root.mainloop()
